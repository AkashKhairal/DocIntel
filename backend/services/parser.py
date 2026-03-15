"""Multi-format document parser supporting PDF, DOCX, Excel, text, and OCR."""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_document(file_path: Path) -> str:
    """Extract text content from a document file.

    Supports: PDF, DOCX, XLSX/XLS/CSV, TXT/MD/JSON/XML, and OCR fallback.
    """
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".pdf":
            return _parse_pdf(file_path)
        elif suffix == ".docx":
            return _parse_docx(file_path)
        elif suffix in (".xlsx", ".xls", ".csv"):
            return _parse_spreadsheet(file_path)
        elif suffix in (".txt", ".md", ".json", ".xml", ".html", ".log", ".py", ".js"):
            return _parse_text(file_path)
        else:
            logger.warning(f"Unsupported file type: {suffix}, attempting text extraction")
            return _parse_text(file_path)
    except Exception as e:
        logger.error(f"Error parsing {file_path}: {e}", exc_info=True)
        return ""


def _parse_pdf(file_path: Path) -> str:
    """Parse PDF using PyMuPDF, with OCR fallback for scanned pages."""
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    text_parts = []

    for page_num, page in enumerate(doc, 1):
        text = page.get_text()
        if text.strip():
            text_parts.append(f"[Page {page_num}]\n{text}")
        else:
            # Attempt OCR on scanned page
            ocr_text = _ocr_page(page)
            if ocr_text:
                text_parts.append(f"[Page {page_num}]\n{ocr_text}")

    doc.close()
    return "\n\n".join(text_parts)


def _ocr_page(page) -> str:
    """OCR a single PDF page using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
        import io

        # Render page to image
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img)
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return ""


def _parse_docx(file_path: Path) -> str:
    """Parse DOCX using python-docx."""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _parse_spreadsheet(file_path: Path) -> str:
    """Parse spreadsheets using pandas."""
    import pandas as pd

    suffix = file_path.suffix.lower()

    try:
        if suffix == ".csv":
            df = pd.read_csv(file_path)
            return df.to_string(index=False)
        else:
            # Excel files — read all sheets
            xls = pd.ExcelFile(file_path)
            parts = []
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                parts.append(f"[Sheet: {sheet_name}]\n{df.to_string(index=False)}")
            return "\n\n".join(parts)
    except Exception as e:
        logger.error(f"Spreadsheet parse error: {e}")
        return ""


def _parse_text(file_path: Path) -> str:
    """Parse plain text files."""
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return file_path.read_bytes().decode("utf-8", errors="replace")
